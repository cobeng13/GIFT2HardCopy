from pathlib import Path
import hashlib
import pytest
from docx import Document
from exam_formatter.docx_engine.exceptions import TemplateContractError
from exam_formatter.docx_engine.generator import generate_exam
from exam_formatter.gift.parser import parse_gift

GIFT = "Question one? {\n=a\n~b\n~c\n~d\n}"

def make_template(path: Path, with_placeholder=True):
    doc = Document()
    for name in ("Exam Question", "Exam Choice", "Exam Choice Compact"):
        doc.styles.add_style(name, 1)
    style = doc.styles["Exam Question"]
    ppr = style.element.get_or_add_pPr(); num = __import__("docx").oxml.OxmlElement("w:numPr"); ppr.append(num)
    doc.add_paragraph("{{EXAM_NAME}}")
    if with_placeholder: doc.add_paragraph("{{QUESTIONS}}")
    doc.save(path)

def test_generation_replaces_metadata_and_preserves_template(tmp_path):
    template, output = tmp_path / "template.docx", tmp_path / "output.docx"; make_template(template)
    digest = hashlib.sha256(template.read_bytes()).hexdigest()
    generate_exam(template, output, {"EXAM_NAME": "Prelims"}, parse_gift(GIFT))
    document = Document(output)
    assert "Prelims" in [p.text for p in document.paragraphs]
    assert sum(p.style.name == "Exam Question" for p in document.paragraphs) == 1
    assert len(document.tables) == 1
    assert all(cell.paragraphs[0].style.name == "Exam Choice Compact" for row in document.tables[0].rows for cell in row.cells)
    assert all(cell.paragraphs[0].paragraph_format.space_after.pt == 6 for cell in document.tables[0].rows[1].cells)
    assert digest == hashlib.sha256(template.read_bytes()).hexdigest()

def test_missing_questions_placeholder(tmp_path):
    template = tmp_path / "template.docx"; make_template(template, False)
    with pytest.raises(TemplateContractError, match="QUESTIONS"):
        generate_exam(template, tmp_path / "output.docx", {}, parse_gift(GIFT))
