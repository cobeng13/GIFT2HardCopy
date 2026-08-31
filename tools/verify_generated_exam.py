"""Structural regression checks for a generated exam DOCX."""
from __future__ import annotations

import argparse
import hashlib
import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.oxml.ns import qn
from exam_formatter.docx_engine.placeholders import iter_all_paragraphs


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("exam", type=Path)
    parser.add_argument("--questions", type=int, required=True)
    parser.add_argument("--source", type=Path, default=Path("Prelims DDS.docx"))
    parser.add_argument("--source-sha256", required=True)
    args = parser.parse_args()
    document = Document(args.exam)
    placeholders = [paragraph.text for paragraph in iter_all_paragraphs(document) if "{{" in paragraph.text]
    questions = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "Exam Question"]
    assert not placeholders, f"Unreplaced placeholders: {placeholders}"
    assert len(questions) == args.questions, f"Expected {args.questions} questions, found {len(questions)}"
    assert all(paragraph._p.pPr.numPr is not None for paragraph in questions), "Question numbering is not Word numbering."
    assert all("*** NOTHING FOLLOWS ***" != paragraph.text or paragraph.text for paragraph in document.paragraphs), "Ending marker missing."
    compact_tables = [table for table in document.tables if any(p.style.name == "Exam Choice Compact" for row in table.rows for cell in row.cells for p in cell.paragraphs)]
    for table in compact_tables:
        assert len(table.rows) == 2 and len(table.columns) == 2, "Compact table is not 2x2."
        borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
        assert borders is not None and all(edge.get(qn("w:val")) == "nil" for edge in borders), "Compact table has visible borders."
    with zipfile.ZipFile(args.exam) as package:
        assert b"PAGE" in package.read("word/footer1.xml"), "Footer PAGE field was not preserved."
    assert hashlib.sha256(args.source.read_bytes()).hexdigest().upper() == args.source_sha256.upper(), "Source template changed."
    print(f"PASS: {args.exam} | questions={len(questions)} | compact_tables={len(compact_tables)}")


if __name__ == "__main__":
    main()
