"""Create the MCQ-only master template from the authoritative source DOCX."""
from __future__ import annotations

import argparse
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def _replace_paragraph_text(paragraph, text: str, run_index: int = 0) -> None:
    paragraph.runs[run_index].text = text
    for run in paragraph.runs[run_index + 1:]:
        run.text = ""


def _add_numbering(style, num_id: int, level: int) -> None:
    p_pr = style.element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(level)); num_pr.append(ilvl)
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id)); num_pr.append(num)
    p_pr.append(num_pr)


def _new_style(document, name: str, source_paragraph, *, question_numbering: bool = False, compact: bool = False):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing or 1.0
    paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    if question_numbering:
        # Reuses the source's real decimal list definition (numId 57, level 0).
        _add_numbering(style, num_id=57, level=0)
        paragraph_format.keep_with_next = True
        paragraph_format.keep_together = True
    elif compact:
        # With the compact table indented to the source label position (0.75 in),
        # these values reproduce the source's 1 in text start after the label.
        paragraph_format.left_indent = Inches(0.25)
        paragraph_format.first_line_indent = Inches(-0.25)
        paragraph_format.space_after = Inches(0)
    else:
        # Source list level 1 uses 1440 twips left / 360 twips hanging.
        paragraph_format.left_indent = Inches(1)
        paragraph_format.first_line_indent = Inches(-0.25)
        paragraph_format.keep_together = True
    return style


def _remove_mcq_body(document, heading, end_marker) -> None:
    body = heading._p.getparent()
    start = body.index(heading._p)
    end = body.index(end_marker._p)
    for element in list(body)[start + 1:end]:
        body.remove(element)
    placeholder = document.add_paragraph(style="Exam Question")
    placeholder.add_run("{{QUESTIONS}}")
    heading._p.addnext(placeholder._p)


def prepare(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    document = Document(destination)
    paragraphs = document.paragraphs
    heading = next(paragraph for paragraph in paragraphs if paragraph.text.startswith("Test 1. Multiple Choice."))
    end_marker = next(paragraph for paragraph in paragraphs if paragraph.text == "*** NOTHING FOLLOWS ***")
    question_source = paragraphs[20]
    choice_source = paragraphs[21]

    _replace_paragraph_text(paragraphs[7], "{{EXAM_NAME}}", run_index=2)
    _replace_paragraph_text(paragraphs[8], "{{COURSE}}")
    _replace_paragraph_text(paragraphs[9], "{{SEMESTER}} ACADEMIC YEAR {{ACADEMIC_YEAR}}")
    _replace_paragraph_text(paragraphs[10], "{{DATE}}")

    _new_style(document, "Exam Question", question_source, question_numbering=True)
    _new_style(document, "Exam Choice", choice_source)
    _new_style(document, "Exam Choice Compact", choice_source, compact=True)
    heading_style = _new_style(document, "Exam Test Heading", heading)
    heading_style.font.bold = True
    heading.style = "Exam Test Heading"
    _remove_mcq_body(document, heading, end_marker)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    prepare(args.source, args.destination)


if __name__ == "__main__":
    main()
