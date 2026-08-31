from __future__ import annotations

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from exam_formatter.config import QUESTION_GAP_AFTER_PT, REQUIRED_STYLES
from exam_formatter.gift.models import Question
from .exceptions import TemplateContractError
from .layout import add_compact_choices, should_use_compact_layout
from .numbering import apply_question_numbering
from .placeholders import find_questions_placeholder, replace_placeholders


def _insert_before(anchor, paragraph) -> None:
    anchor._p.addprevious(paragraph._p)


def _require_styles(document) -> None:
    available = {style.name for style in document.styles}
    for name in REQUIRED_STYLES:
        if name not in available:
            raise TemplateContractError(f"Required Word style '{name}' was not found.")


def answer_key_path(output: Path) -> Path:
    output = Path(output)
    return output.with_name(f"{output.stem}_ANSWERKEY.txt")


def _answer_key_text(questions: list[Question]) -> str:
    lines = ["ANSWER KEY", ""]
    for number, question in enumerate(questions, 1):
        correct_index = next(index for index, choice in enumerate(question.choices) if choice.correct)
        lines.append(f"{number}. {chr(65 + correct_index)}")
    return "\n".join(lines) + "\n"


def generate_exam(template: Path, output: Path, metadata: dict[str, str], questions: list[Question]) -> Path:
    template = Path(template)
    output = Path(output)
    if not template.is_file():
        raise FileNotFoundError(f"Master template was not found: {template}")
    if template.resolve() == output.resolve():
        raise ValueError("Output file must be different from the master template.")
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.tmp.docx")
    key_output = answer_key_path(output)
    temporary_key = key_output.with_name(f".{key_output.stem}.tmp.txt")
    try:
        shutil.copy2(template, temporary)
        document = Document(temporary)
        _require_styles(document)
        replace_placeholders(document, {f"{{{{{key}}}}}": value for key, value in metadata.items()})
        placeholder = find_questions_placeholder(document)
        if placeholder is None:
            raise TemplateContractError("Master template does not contain {{QUESTIONS}} as a body paragraph.")
        for question in questions:
            stem = document.add_paragraph(style="Exam Question")
            stem.add_run(question.stem)
            stem.paragraph_format.keep_with_next = True
            stem.paragraph_format.keep_together = True
            stem.paragraph_format.widow_control = True
            apply_question_numbering(stem)
            _insert_before(placeholder, stem)
            if should_use_compact_layout(question):
                add_compact_choices(document, placeholder._p, question)
            else:
                for index, choice in enumerate(question.choices):
                    choice_paragraph = document.add_paragraph(style="Exam Choice")
                    choice_paragraph.add_run(f"{chr(65 + index)}. {choice.text}")
                    choice_paragraph.paragraph_format.keep_together = True
                    if index == len(question.choices) - 1:
                        choice_paragraph.paragraph_format.space_after = Pt(QUESTION_GAP_AFTER_PT)
                    _insert_before(placeholder, choice_paragraph)
        placeholder._element.getparent().remove(placeholder._element)
        document.save(temporary)
        temporary_key.write_text(_answer_key_text(questions), encoding="utf-8")
        temporary.replace(output)
        temporary_key.replace(key_output)
        return key_output
    except Exception:
        if temporary.exists():
            temporary.unlink()
        if temporary_key.exists():
            temporary_key.unlink()
        raise
