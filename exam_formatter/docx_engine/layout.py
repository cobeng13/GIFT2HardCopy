from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from exam_formatter.config import QUESTION_GAP_AFTER_PT, TWO_COLUMN_MAX_CHARS
from exam_formatter.gift.models import Question


def should_use_compact_layout(question: Question) -> bool:
    return all("\n" not in choice.text and len(choice.text) <= TWO_COLUMN_MAX_CHARS for choice in question.choices)


def set_borderless(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    properties = table._tbl.tblPr
    old = properties.first_child_found_in("w:tblBorders")
    if old is not None:
        properties.remove(old)
    properties.append(borders)


def set_table_indent(table, twips: int) -> None:
    properties = table._tbl.tblPr
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(twips))
    indent.set(qn("w:type"), "dxa")
    old = properties.first_child_found_in("w:tblInd")
    if old is not None:
        properties.remove(old)
    properties.append(indent)


def add_compact_choices(document, anchor, question: Question) -> None:
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_borderless(table)
    # Word's default left cell margin is 108 twips. Offset the table by that
    # amount so compact A/B labels start at the same 0.75 in position as the
    # source's level-1 vertical choices.
    set_table_indent(table, 972)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Inches(3.1)
    values = ((0, 2), (1, 3))
    for row_index, indexes in enumerate(values):
        for column_index, choice_index in enumerate(indexes):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            paragraph.style = "Exam Choice Compact"
            paragraph.add_run(f"{chr(65 + choice_index)}. {question.choices[choice_index].text}")
            if row_index == 1:
                paragraph.paragraph_format.space_after = Pt(QUESTION_GAP_AFTER_PT)
    anchor.addprevious(table._tbl)
