from __future__ import annotations

from collections.abc import Mapping
from docx.document import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def _paragraphs_in_table(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _paragraphs_in_table(nested)


def iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _paragraphs_in_table(table)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                yield from _paragraphs_in_table(table)


def replace_in_paragraph(paragraph: Paragraph, replacements: Mapping[str, str]) -> bool:
    original = "".join(run.text for run in paragraph.runs)
    replacement = original
    for token, value in replacements.items():
        replacement = replacement.replace(token, value)
    if replacement == original:
        return False
    if not paragraph.runs:
        paragraph.add_run(replacement)
    else:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def replace_placeholders(document: Document, replacements: Mapping[str, str]) -> None:
    for paragraph in iter_all_paragraphs(document):
        replace_in_paragraph(paragraph, replacements)


def find_questions_placeholder(document: Document) -> Paragraph | None:
    return next((paragraph for paragraph in document.paragraphs if "{{QUESTIONS}}" in "".join(run.text for run in paragraph.runs)), None)
